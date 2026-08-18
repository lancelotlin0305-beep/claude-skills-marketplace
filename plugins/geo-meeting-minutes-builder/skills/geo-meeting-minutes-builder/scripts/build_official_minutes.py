# -*- coding: utf-8 -*-
"""公文體正式會議紀錄產生器:JSON → .docx(以 assets/minutes-official-style.docx 為樣式基底)

用法: python build_official_minutes.py <content.json> <輸出.docx>

JSON schema(欄位皆字串,陣列如列;來源標註/時間戳不放進正式版):
{
  "title": "法務部 115 年度「…」第二次…討論會議",
  "subject": "會議主題一句話",
  "time": "115/08/14 上午 10 時",            # 民國紀年
  "location": "…會議室 / 線上",
  "attendees": "詳如簽到表",                  # 或逐一列名
  "proposals": [
    { "title": "書類製作系統轉登…之情境與資料拋轉需求。",
      "explain": ["說明條列一", "…"],
      "resolutions": ["決議條列一", "…"] }
  ],
  "motions": "無。",                          # 臨時動議;有內容則為條列陣列亦可
  "tracking": [
    { "item": "追蹤事項描述", "owner": "負責單位", "due": "115/08/21|未定",
      "status": "進行中|已完成" }
  ],
  "chair": "會議主持人姓名職稱",
  "recorder": "會議記錄人姓名"
}
"""
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

CN_NUM = "一二三四五六七八九十"


def cn_index(i):
    """1-based → 中文數字(支援至 19,會議提案/追蹤足矣)。"""
    if i <= 10:
        return CN_NUM[i - 1]
    return "十" + CN_NUM[i - 11]


def set_cjk(run, size=None, bold=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def para(doc, text, size=13, bold=False, align=None, indent=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cjk(r, size, bold)
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    p.paragraph_format.space_after = Pt(4)
    return p


def main():
    if len(sys.argv) != 3:
        sys.exit(__doc__)
    m = json.loads(Path(sys.argv[1]).read_text(encoding="utf-8"))

    # 以範本為樣式基底(頁面設定/預設樣式),清空本文重建
    tpl = Path(__file__).parent.parent / "assets" / "minutes-official-style.docx"
    doc = Document(str(tpl))
    body = doc.element.body
    for el in list(body):
        if el.tag.endswith("}sectPr"):
            continue
        body.remove(el)

    # 標題
    para(doc, m["title"], size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "會議紀錄", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 表頭
    para(doc, f"會議主題:{m['subject']}")
    para(doc, f"會議時間:{m['time']}")
    para(doc, f"會議地點:{m['location']}")
    para(doc, f"與會人員:{m.get('attendees', '詳如簽到表')}")
    para(doc, "討論議題:")

    # 提案
    for i, pr in enumerate(m.get("proposals", []), 1):
        para(doc, f"提案{cn_index(i)}、{pr['title']}", bold=True)
        para(doc, "說明:")
        for j, s in enumerate(pr.get("explain", []), 1):
            para(doc, f"{cn_index(j)}、{s}", indent=24)
        para(doc, "決議:")
        for j, s in enumerate(pr.get("resolutions", []), 1):
            para(doc, f"{cn_index(j)}、{s}", indent=24)

    # 臨時動議
    motions = m.get("motions", "無。")
    if isinstance(motions, list):
        para(doc, "臨時動議:", bold=True)
        for j, s in enumerate(motions, 1):
            para(doc, f"{cn_index(j)}、{s}", indent=24)
    else:
        para(doc, f"臨時動議:{motions}", bold=True)
    para(doc, "散會。")

    # 追蹤事項表
    tracking = m.get("tracking", [])
    if tracking:
        para(doc, "追蹤事項:", bold=True)
        tb = doc.add_table(rows=1 + len(tracking), cols=5)
        tb.style = "Table Grid"
        heads = ["項次", "追蹤事項", "負責單位", "完成期限", "狀態"]
        for c, h in enumerate(heads):
            r = tb.rows[0].cells[c].paragraphs[0].add_run(h)
            set_cjk(r, 12, True)
        for ri, t in enumerate(tracking, 1):
            vals = [cn_index(ri), t["item"], t.get("owner", ""),
                    t.get("due", "未定"), t.get("status", "進行中")]
            for c, v in enumerate(vals):
                r = tb.rows[ri].cells[c].paragraphs[0].add_run(str(v))
                set_cjk(r, 12)

    # 落款
    para(doc, "")
    para(doc, f"會議主持人 : {m.get('chair', '')}")
    para(doc, f"會議記錄人 : {m.get('recorder', '')}")

    doc.save(sys.argv[2])
    print(f"寫出:{sys.argv[2]}")


if __name__ == "__main__":
    main()
